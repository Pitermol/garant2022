import json
import os
import re
import configparser

import docx
import jinja2
from dataclasses import dataclass
import requests
import fitz
from typing import List, Optional
import logging
import odf
import ezodf
import os, sys
import zipfile
import xml.dom.minidom
import chardet
from odf.opendocument import load


@dataclass
class Document:  # Создается класс для удобного хранения данных о документе в виде объекта
    num: str  # Номер по счету
    index: str  # $topic
    link: str  # Ссылка
    name: str  # Названи документа
    isActiveThen: bool  # Действителен ли на данную дату
    isActiveNow: bool  # Действителен ли на сегодня
    context: list  # Контекст
    date: str  # Дата для "isActiveThen"


class checkDocsRelevance:  # Основной класс с методами для формирования списка и HTML - таблицы
    def __init__(self, name, date, new_name):
        self.dateToCheck = date
        self.docxFile = None
        self.new_name = new_name
        self.docRawText = ""
        self.docCreatingDate = ""
        self.fileName = name
        self.operationStatus = True
        self.docType = self.fileName.split(".")[-1]
        config = configparser.ConfigParser()
        config.read('settings.ini')
        self.APIToken = config["api"]["token"]
        self.apiUrlMakeLinks = config["api"]["url_make_links"]
        self.apiAccept = config["api"]["accept"]
        self.apiContentType = config["api"]["content-type"]
        self.linksBaseUrl = config["api"]["linksBaseUrl"]
        self.apiUrlGetModifications = config["api"]["apiUrlGetModifications"]
        self.apiUrlGetDocInfo = config["api"]["apiUrlGetDocInfo"]

    def open_doc(self):  # Функция для открытия документа и его считывания
        try:
            if self.docType == "odt":
                doc = ezodf.opendoc(f"input/{self.fileName}")
                print(doc.meta)
                m_odf = zipfile.ZipFile(f"input/{self.fileName}")
                filelist = m_odf.infolist()
                ostr = m_odf.read('content.xml')
                meta = m_odf.read("META-INF/manifest.xml")
                print(meta)
                doc = xml.dom.minidom.parseString(ostr)
                paras = doc.getElementsByTagName('text:p')

                text_in_paras = []
                for p in paras:
                    for ch in p.childNodes:
                        if ch.nodeType == ch.TEXT_NODE:
                            text_in_paras.append(ch.data)
                self.docRawText = "".join(text_in_paras)

                self.docRawText = self.docRawText.replace("\n\n", "\n")
                self.docRawText = self.docRawText.replace("\t", " ")
                self.docRawText = self.docRawText.strip()
                while "  " in self.docRawText:
                    self.docRawText = self.docRawText.replace("  ", " ")
            elif self.docType == "pdf":  # Если формат - PDF
                with fitz.open(f'input/{self.fileName}') as doc:
                    metadata = doc.metadata["creationDate"]
                    if len(metadata) > 0:
                        year = metadata[2:][:4]
                        month = metadata[2:][4:6]
                        day = metadata[2:][6:8]
                        self.docCreatingDate = f"{day}.{month}.{year}"
                    text = ""
                    for page in doc:
                        text += page.get_text()
                self.docRawText = text

                self.docRawText = self.docRawText.replace("\n\n", "\n")
                self.docRawText = self.docRawText.replace("\t", " ")
                self.docRawText = self.docRawText.strip()
                while "  " in self.docRawText:
                    self.docRawText = self.docRawText.replace("  ", " ")
            elif self.docType == "docx":  # Если формат - DOCX
                self.docxFile = docx.Document(f'input/{self.fileName}')
                self.docCreatingDate = self.docxFile.core_properties.created
                for para in range(len(self.docxFile.paragraphs)):
                    self.docRawText += " " + self.docxFile.paragraphs[para].text

                self.docRawText = self.docRawText.replace("\n\n", "\n")
                self.docRawText = self.docRawText.replace("\t", " ")
                self.docRawText = self.docRawText.strip()
                while "  " in self.docRawText:
                    self.docRawText = self.docRawText.replace("  ", " ")
            elif self.docType == "txt":  # Если формат - TXT
                self.docxFile = open(f'input/{self.fileName}', "r")
                self.docRawText = self.docxFile.read()

                self.docRawText = self.docRawText.replace("\n\n", "\n")
                self.docRawText = self.docRawText.replace("\t", " ")
                self.docRawText = self.docRawText.strip()
                while "  " in self.docRawText:
                    self.docRawText = self.docRawText.replace("  ", " ")
        except Exception as e:
            logging.error(e)

            result = {"successStatus": "False", "errorCode": "1", "sourceFileName": self.fileName}
            with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
                f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
            f.close()
            self.operationStatus = False

            return

    def get_text(self) -> str:  # Возвращает полученный из документа текст
        return self.docRawText

    def make_doc_list(self) -> Optional[List[Document]]:  # Метод для создания списка из экземпляров класса "Document"
        self.open_doc()  # Получение текста на вход
        requestslist = []
        docsLinksList = []
        url = self.apiUrlMakeLinks  # Ссылка для запроса
        # print(url)
        # print(self.docCreatingDate)
        headers = {  # Заголовки для запроса
            'Accept': self.apiAccept,
            'Content-type': self.apiContentType,
            'Authorization': f'Bearer {self.APIToken}'
        }

        # print(self.docRawText)
        text = ""
        ind = 0
        paraList = self.docRawText.split("\n")
        print(len(paraList))
        cur_text = ""
        for i, paragraph in enumerate(paraList):
            if len(cur_text) + len(paragraph) <= 2000:
                cur_text += paragraph
                if i + 1 != len(paraList):
                    continue

            if cur_text == "":
                cur_text = paragraph
            # print(len(cur_text))
            curSentences = ""
            for j, sentence in enumerate(re.split(",", cur_text)):
                # print(sentence)
                if len(curSentences) + len(sentence) <= 2000:
                    curSentences += sentence + ","
                    if j + 1 != len(re.split(",", cur_text)):
                        continue
                cur = curSentences[:-1]
                print(cur)
                payload = json.dumps({
                    "text": cur,
                    "baseUrl": self.linksBaseUrl
                })
                try:
                    response = requests.request("POST", url, headers=headers,
                                            data=payload, timeout=30)  # Запрос для проставления ссылок в отрывке текста
                    print(response.text)
                except Exception as e:
                    logging.error(str(e) + ": Не удалось отправить запрос для проставления ссылок")
                    result = {"successStatus": "False", "errorCode": "2", "sourceFileName": self.fileName}
                    with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
                        f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
                    f.close()
                    self.operationStatus = False
                    return []
                html = response.json()["text"].replace("&quot;", "\"").replace("ГАРАНТ1/2", "").replace("20072022", " ")
                if ind != 2000:
                    center = text[-30:] + html[:30]
                    if "<a" not in center and "a>" not in center and "</a" not in center:
                        payload = json.dumps({
                            "text": center,
                            "baseUrl": self.linksBaseUrl
                        })
                        try:
                            response = requests.request("POST", url, headers=headers,
                                                    data=payload,
                                                    timeout=30)  # Запрос для проставления ссылок в отрывке текста
                            text = text[:-30] + response.json()["text"] + html[30:]
                        except Exception as e:
                            logging.error(str(e) + ": Не удалось отправить запрос для проставления ссылок")
                            result = {"successStatus": "False", "errorCode": "3", "sourceFileName": self.fileName}
                            with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
                                f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
                            f.close()
                            self.operationStatus = False
                            return []
                    else:
                        text += html
                else:
                    text += html
                curSentences = ""
            cur_text = ""

        html = text
        print(html)
        htmls = re.split("<a href=\"|</a>", html)
        # print(htmls)
        # print(html)
        ind = 0

        for i in range(1, len(htmls),
                       2):  # Поиск всех документов, извлечение ссылок, их номеров в системе Гарант и контекста
            ind += 1
            htmlc = htmls[i]
            htmlc = re.split('"', htmlc)
            link = htmlc[0]  # Ссылка
            docLength = len(htmls[i].split(">")[1])
            # print(docLength, htmls[i-2])
            curDocContext = ""
            htmlPart = i - 1

            while htmlPart >= 0 and len(curDocContext) < (244 - docLength) // 2:
                if "https" in htmls[htmlPart]:
                    partNeeded = htmls[htmlPart].split(">")[1]
                    if len(partNeeded) + len(curDocContext) <= (244 - docLength) // 2:
                        curDocContext = partNeeded + curDocContext
                    else:
                        # print(1.2, len(partNeeded) + len(curDocContext), (244 - docLength) // 2)
                        # curDocContext = partNeeded[-((244 - docLength) // 2 - len(curDocContext)) + 1:] + curDocContext
                        curDocContext = partNeeded + curDocContext
                        curDocContext = curDocContext[-((244 - docLength) // 2):]
                        # print(len(curDocContext))
                    htmlPart -= 1
                    # print(1, curDocContext)
                else:
                    if len(htmls[htmlPart]) + len(curDocContext) <= (244 - docLength) // 2:
                        curDocContext = htmls[htmlPart] + curDocContext
                        # print(2.1, curDocContext)
                    else:
                        # curDocContext = htmls[htmlPart][-((244 - docLength) // 2 - len(curDocContext)) + 1:] + curDocContext
                        curDocContext = htmls[htmlPart] + curDocContext
                        curDocContext = curDocContext[-((244 - docLength) // 2):]
                        # print(2.2, curDocContext)
                    htmlPart -= 1

            # print(len(curDocContext), curDocContext)
            # print()

            linkLength = len(htmls[i].split("\"")[0])
            curDocContext += "<a href=\"" + htmls[i] + "</a>"

            htmlPart = i + 1
            while htmlPart < len(htmls) and len(curDocContext) < 250 + 7 + linkLength:
                if "https" in htmls[htmlPart]:
                    partNeeded = htmls[htmlPart].split(">")[1]
                    # print(partNeeded)
                    if len(partNeeded) <= 250 + 7 + linkLength - len(curDocContext):
                        curDocContext += partNeeded
                    else:
                        curDocContext += partNeeded
                        curDocContext = curDocContext[:250 + 7 + linkLength]
                        # curDocContext += partNeeded[:250 + docLength + 15 - len(curDocContext)]
                    htmlPart += 1
                else:
                    if len(htmls[htmlPart]) <= 250 + linkLength + 7 - len(curDocContext):
                        curDocContext += htmls[htmlPart]
                    else:
                        curDocContext += htmls[htmlPart]
                        curDocContext = curDocContext[:250 + linkLength + 7]
                        # curDocContext += htmls[htmlPart][:250 + docLength + 15 - len(curDocContext)]
                    htmlPart += 1

            # try:
            #     if len(htmls) > i + 1:
            #         curDocContext += htmls[i + 1][:(244 - docLength) // 2]
            # except Exception as e:
            #     print(e)
            #     if len(htmls) > i + 1:
            #         curDocContext += htmls[i + 1]

            curDocContext = curDocContext.replace("<p>", "")
            curDocContext = curDocContext.replace("</p>", "")  # Контекст
            # print(len(curDocContext), curDocContext)

            if link not in docsLinksList:
                docsLinksList.append(link)
                curDocContext = ["..." + curDocContext + "..."]
            else:
                requestslist[requestslist.index(list(filter(lambda x: x.link == link, requestslist))[0])].context.append(f" ...{curDocContext}...")
                ind -= 1
                # print(1)
                continue

            htmlc = str(htmlc[0])
            htmlc = htmlc.split('document/')
            htmlc = htmlc[1]
            htmlc = htmlc.split('/')
            number = htmlc[0]  # Номер в системе Гарант

            url = self.apiUrlGetModifications

            payload = json.dumps({"topics": [number], "modDate": self.dateToCheck})
            headers = {
                'Authorization': f'Bearer {self.APIToken}',
                'Content-type': self.apiContentType,
                'Accept': self.apiAccept
            }
            try:
                response = requests.post(url, headers=headers, data=payload, timeout=30).json()[
                    "topics"]  # Запрос для получения сведений об акутальности документа на данную дату
            except Exception as e:
                logging.error(str(e) + ": Не удалось отправить запрос для получения изменений в документе с данной даты")

                result = {"successStatus": "False", "errorCode": "6", "sourceFileName": self.fileName}
                with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
                    f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
                f.close()
                self.operationStatus = False

                return []

            if len(response) == 0:
                isActiveThen = True
            else:
                isActiveThen = False
            # print(number)
            url1 = f'{self.apiUrlGetDocInfo}{number}'
            headers1 = {
                'Accept': self.apiAccept,
                'Content-type': self.apiContentType,
                'Authorization': f'Bearer {self.APIToken}',
            }

            try:
                response1 = requests.get(url1,
                                     headers=headers1, timeout=30)  # Запрос для получения названия документа и сведений об актуальности на сегодня
                print(response1)
                response1 = response1.json()
            except Exception as e:
                logging.error(str(e) + ": Не удалось отправить запрос для получения информации о документе")

                result = {"successStatus": "False", "errorCode": "7", "sourceFileName": self.fileName}
                with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
                    f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
                f.close()
                self.operationStatus = False

                return []
            # print(response1)
            curDocName = response1['name']
            if len(curDocName) > 250:
                curDocName = curDocName[:250]
            isActiveNow = response1['status']
            if isActiveNow == "Действующие":
                isActiveNow = True
            else:
                isActiveNow = False

            a = Document(str(ind), number, link, curDocName, isActiveThen, isActiveNow, curDocContext,
                         self.dateToCheck)  # Создание объекта документа
            requestslist.append(a)

        return requestslist

    def create_table(self):  # Метод для формирования HTML - Файла с таблицой документов
        data = self.make_doc_list()
        if not self.operationStatus:
            return

        changed = inactive = 0
        docLinksList = []
        for i, doc in enumerate(data, 1):
            docLinksList.append({"id": i, "docName": doc.name, "docLink": doc.link,
                                 "abolishedStatus": "True" if not doc.isActiveNow else "False",
                                 "changedStatus": "True" if not doc.isActiveThen else "False",
                                 "linkContext": doc.context})
            if not doc.isActiveNow:
                inactive += 1
            if not doc.isActiveThen:
                changed += 1

        # print(data)

        template = jinja2.Template('''<!DOCTYPE html>
        <html>
            <head>
                <title>Docs list</title>
                <link rel="stylesheet" href="style.css">
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.2.0-beta1/dist/css/bootstrap.min.css"  rel="stylesheet">
                <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.2.0-beta1/dist/js/bootstrap.bundle.min.js"></script>
            </head>
            <body>
                <header></header>
                <div style="margin-top:3%" align="center">
                    <table class="table table-striped" style="width:90%; word-break: normal;">
                        <thead>
                            <tr>
                                <th scope="col">#</th>
                                <th scope="col">Документ</th>
                                <th scope="col">Статус на сегодня</th>
                                <th scope="col">Изменялся ли с {{given_date}}</th>
                                <th scope="col">Контекст</th>
                            </tr>
                        </thead>
                        <tbody>{% for doc in docs %}
                            <tr>
                                <th scope="row">{{doc.num}}</th>
                                <td>{{doc.name}}</td>{% if doc.isActiveNow %}
                                <td>Действующий</td>{% else %}
                                <td>Недействующий</td>{% endif %}{% if doc.isActiveThen %}
                                <td>Нет</td>{% else %}
                                <td>Да</td>{% endif %}
                                <td>{% for context in doc.context %}
                                    {{context}}<br>{% endfor %}
                                </td>
                            </tr>{% endfor %}
                        </tbody>
                    </table>
                </div>
            </body>
        </html>''')
        try:
            html = template.render(docs=data, given_date=self.dateToCheck)  # Формирование HTML
        except Exception as e:
            logging.error(str(e) + ": Не удалось сформировать HTML")

            result = {"successStatus": "False", "errorCode": "8", "sourceFileName": self.fileName}
            with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
                f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
            f.close()

            return

        result = {"successStatus": "True", "errorCode": None, "sourceFileName": self.fileName,
                  "docLinksCount": len(data), "abolishedDocsCount": inactive,
                  "changedDocsCount": changed, "htmlResult": f"{self.new_name}.html", "docLinksList": docLinksList}

        with open(f"output/{self.new_name}.json", "w", encoding="utf-8") as f:
            f.write(json.dumps(result, ensure_ascii=False, indent=4, sort_keys=False))
        f.close()

        # print(html)
        with open(f"output/{self.new_name}.html", "w", encoding='utf-8') as f:  # Запись в файл "output/template.html"
            f.write(html)
        f.close()


def start(date="2021-10-25"):  # Функция обработки множества файлов в папке input
    input_files = os.listdir("input")  # Список входных файлов
    for old_file in os.listdir("output"):
        os.remove(os.path.join(os.path.abspath(os.path.dirname(__file__)) + "/output", old_file))  # Удаление всех старых файлов в папке output

    for i, newFile in enumerate(input_files):  # Перебор входных файлов
        obj = checkDocsRelevance(newFile, date, f"out_{str(i + 1)}")  # Создание экземпляра основного класса
        obj.create_table()  # Вызов метода для формирования таблицы


start()  # 250 символов в контексте, дату из сведений
